""" """

import os
import docx

import db.models as models

class DocumentParser():
    def __init__(self, filename):
        if not os.path.exists(filename := os.path.abspath(filename)):
            raise docx.opc.exceptions.PackageNotFoundError("Invalid path to the file with the document")

        self.filename = filename
        self.document = docx.Document(self.filename)

    def discipline_base(self):
        iterator = BaseInfoIterator(self.document)
        fields = [None] + [field for field in iterator]
        return models.DisciplineWorkProgram(*fields)

    def learning_outcomes(self):
        pass

    def educational_program(self):
        iterator = EducationalProgramIterator(self.document)
        models_list = list()
        for educational_program in iterator:
            models_list.append(models.EducationalProgram(*educational_program))

        return models_list

    def discipline_scope(self):
        iterator = SemesterScopeIterator(self.document)
        raw_objects = list()
        for i, semester_scope in enumerate(iterator):
            if i <= 1:
                raw_objects.append([None, 0] + semester_scope)
            else:
                for j, data in enumerate(semester_scope):
                    raw_objects[j].append(data)

        models_list = list()
        for obj in raw_objects:
            models_list.append(models.DisciplineScope(*obj))

        return models_list

    def discipline_module(self):
        pass

    def discipline_materials(self):
        iterator = MaterialIterator(self.document)
        models_list = list()
        for materials in iterator:
            models_list.append(models.DisciplineMaterial(None, 0, materials))

        return models_list

    def get_discipline_program(self):
        base_model = self.discipline_base()
        base_model.educational_program = self.educational_program()
        base_model.discipline_material = self.discipline_materials()
        base_model.discipline_scope = self.discipline_scope()
        return base_model

    def save(self):
        self.document.save()


class BaseInfoIterator():
    def __iter__(self):
        return self

    def __init__(self, document):
        self.document = document

        self.counter = 0
        self.trigger_index = 0
        self.triggers = (
            lambda x: "рабочая программа дисциплины" not in x,
            lambda x: "автор программы:" not in x,
            lambda x: not x.endswith("+"),
        )

        self.offsets = (lambda: +1, lambda: +1, lambda: +0)
        self.formatter = (lambda x: x[1:-1], lambda x: x[:x.find(",")], lambda x: x)

    def __next__(self):
        if self.trigger_index >= len(self.triggers):
            raise StopIteration

        paragraphs = self.document.paragraphs
        current_trigger = self.triggers[self.trigger_index]
        current_offset = self.offsets[self.trigger_index]
        current_formatter = self.formatter[self.trigger_index]

        while current_trigger(paragraphs[self.counter].text.lower()):
            self.counter += 1

        self.trigger_index += 1
        return current_formatter(paragraphs[self.counter + current_offset()].text)


class EducationalProgramIterator():
    def __iter__(self):
        return self

    def __init__(self, document):
        self.document = document
        self.counter = 0
        self.formatter = lambda x: x[x.find("по направлениям ") + len("по направлениям "):-1]

        paragraphs = self.document.paragraphs
        index = 0
        while "дисциплина входит" not in paragraphs[index].text.lower():
            index += 1

        self.raw_pairs = self.formatter(paragraphs[index].text).split(',')

    def __next__(self):
        if self.counter >= len(self.raw_pairs):
            raise StopIteration

        program_id, name = self.raw_pairs[self.counter].strip().split(' ', 1)
        self.counter += 1
        return program_id, name[1:-1]


class MaterialIterator():
    def __iter__(self):
        return self

    def __init__(self, document):
        self.document = document
        self.counter = 0
        self.formatter = lambda x: x[x.find("по направлениям ") + len("по направлениям "):-1]
        self.start_trigger = "основная литература по дисциплине"
        self.end_trigger = "перечень ресурсов сети интернет, "
        self.skip_trigger = "дополнительные учебные материалы"
        self.paragraphs = self.document.paragraphs

        while self.start_trigger not in self.paragraphs[self.counter].text.lower():
            self.counter += 1

    def __next__(self):
        self.counter += 1
        current_material = self.paragraphs[self.counter].text.lower().strip()
        if self.end_trigger in current_material:
            raise StopIteration

        if self.skip_trigger in current_material:
            self.counter += 1
            current_material = self.paragraphs[self.counter].text.lower().strip()

        while current_material == "":
            self.counter += 1
            current_material = self.paragraphs[self.counter].text.lower().strip()

        if self.end_trigger in current_material:
            raise StopIteration

        return self.paragraphs[self.counter].text.strip()


class SemesterScopeIterator():
    def __iter__(self):
        return self

    def __init__(self, document):
        self.document = document
        self.table = self.document.tables[3]
        self.paragraphs = self.document.paragraphs

        self.counter = 0
        self.table_trigger_index = 0
        self.text_trigger_index = 0
        self.table_triggers = (
            lambda x: "объем дисциплины" in x,
            lambda x: "лекции" in x,
            lambda x: "семинары" in x,
            lambda x: "лабораторные работы" in x,
            lambda x: "самостоятельная работа" in x,
            lambda x: "вид промежуточной аттестации" in x,
        )

        self.text_trigger = lambda x: "объем дисциплины составляет" not in x

        while self.text_trigger(self.paragraphs[self.counter].text.lower()):
            self.counter += 1

        self.counter +=2

    def parse_table(self, row_ind, cell_ind):
        current_trigger = self.table_triggers[self.table_trigger_index]
        self.table_trigger_index +=1

        for i, row in enumerate(self.table.rows):
            for j, cell in enumerate(row.cells):
                if current_trigger(cell.text.lower()):
                    row_ind = i
                    cell_ind = j + 2
                    break

        cells = self.table.rows[row_ind].cells
        cell = cells[cell_ind].text
        cells_list = list()

        while cell.replace('.', '', 1).isdigit() or cell.lower() in ('экзамен', 'зачёт', 'зачет'):
            try:
                cells_list.append(round(float(cell)))
            except ValueError:
                cells_list.append(cell)

            cell_ind += 1
            if cell_ind >= len(cells):
                return cells_list

            cell = cells[cell_ind].text

        return cells_list

    def __next__(self):
        if self.table_trigger_index >= len(self.table_triggers):
            raise StopIteration

        if self.text_trigger_index <= 1:
            current = self.paragraphs[self.counter].text.lower().strip()
            self.text_trigger_index += 1

            while current == "":
                current = self.paragraphs[self.counter].text.lower()
                self.counter += 1

            semester_number, raw = current.split(' ', 1)
            for value in raw:
                if value.replace('.', '', 1).isdigit():
                    self.counter += 1
                    return [semester_number, round(float(value))]

        return self.parse_table(0, 0)

